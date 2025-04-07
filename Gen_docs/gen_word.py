from docx import Document

# Crear un nuevo documento Word
doc = Document()

# Agregar título y contenido
doc.add_heading('Mi primer archivo Word con Python', 0)
doc.add_paragraph('Hola, este archivo fue generado automáticamente.')

# Agregar una lista
doc.add_paragraph('Elemento 1', style='List Bullet')
doc.add_paragraph('Elemento 2', style='List Bullet')

# Guardar el archivo
doc.save('mi_archivo.docx')

print("Documento Word creado exitosamente.")
